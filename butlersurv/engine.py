from .models import QuestionSurv
import docx
from datetime import datetime

def add_answer(document, spisok):
    if (len(spisok) > 0) and (spisok[0] != ''):
        for element in spisok:
            document.add_paragraph(element, 'List Bullet 2')

def sent_to_file(data):
    sys = []
    part = []
    file_name = str(datetime.now()) + '.docx'
    document = docx.Document()
    document.add_heading('ДЕТАЛЬНЫЙ ОПРОСНЫЙ ЛИСТ', 0)
    for key in data.keys():
        query_bd = QuestionSurv.objects.get(id=key)
        if (query_bd.partition.type.id in sys) and (query_bd.partition.id in part):
                p = document.add_paragraph('','List Bullet')
                r = p.add_run(query_bd.question)
                r.bold = True
                add_answer(document, data[key])
        elif query_bd.partition.type.id in sys:
                document.add_paragraph(query_bd.partition.name, 'Intense Quote')
                p = document.add_paragraph('', 'List Bullet')
                r = p.add_run(query_bd.question)
                r.bold = True
                add_answer(document, data[key])
                part.append(query_bd.partition.id)
        else:
            document.add_heading(query_bd.partition.type.name, 1)
            document.add_paragraph(query_bd.partition.name, 'Intense Quote')
            p = document.add_paragraph('','List Bullet')
            r = p.add_run(query_bd.question)
            r.bold = True
            add_answer(document, data[key])
            sys.append(query_bd.partition.type.id)
            part.append(query_bd.partition.id)

    document.save('./media/surv/'+ file_name)
    return file_name